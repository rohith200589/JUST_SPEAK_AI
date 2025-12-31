# Import necessary libraries
from flask import Flask, jsonify
from flask_cors import CORS
from flask_graphql import GraphQLView
import graphene
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound, VideoUnavailable, TranscriptsDisabled
from urllib.parse import urlparse, parse_qs
import yt_dlp
import os
import tempfile
import shutil # Import shutil for rmtree
import ctranslate2  # Make sure this is installed via: pip install ctranslate2
import time # Import time for measuring transcription duration
import base64 # For handling base64 encoded file content
import requests # Re-import requests for Gemini API
import json # Re-import json for Gemini API payload
import re
# Import Flask-SocketIO
from flask_socketio import SocketIO, emit # Import SocketIO and emit


# Imports for Faster Whisper and audio processing
from faster_whisper import WhisperModel
from pydub import AudioSegment
import math
try:
    # When imported as a package (e.g., tests), use relative import
    from .gemini_utils import _strip_code_fences_and_markdown, _extract_first_json, _save_raw_gemini_response
except Exception:
    # When running `python server.py` directly, fall back to direct import
    from gemini_utils import _strip_code_fences_and_markdown, _extract_first_json, _save_raw_gemini_response


# Small custom exception to signal Gemini API conditions (e.g., quota exceeded)
class GeminiAPIError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code

# Imports for document parsing
import PyPDF2 # pip install PyPDF2
from docx import Document # pip install python-docx


app = Flask(__name__)
CORS(app)

socketio = SocketIO(app, cors_allowed_origins=["*","https://just-speak-nine.vercel.app"], async_mode='gevent')

@app.route('/')
def index():
    return "Flask backend is running!"

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "ok", "service": "main_server"}), 200


# Define the GraphQL types
class Timestamp(graphene.ObjectType):
    time = graphene.Float()
    text = graphene.String()

# Define the new GraphQL types for analysis
class Metric(graphene.ObjectType):
    title = graphene.String()
    value = graphene.String()
    change = graphene.String()
    changeType = graphene.String()

class RephraseSuggestion(graphene.ObjectType):
    original = graphene.String()
    suggested = graphene.String()

class EngagementData(graphene.ObjectType):
    segment = graphene.String()
    engagement = graphene.Int()
    
class AnalysisResult(graphene.ObjectType):
    metrics = graphene.List(Metric)
    rephraseSuggestions = graphene.List(RephraseSuggestion)
    engagementChartData = graphene.List(EngagementData)

class KeyInsightsResult(graphene.ObjectType):
    title = graphene.String()
    points = graphene.List(graphene.String)

class SummaryResult(graphene.ObjectType):
    title = graphene.String()
    subheader = graphene.String()
    points = graphene.List(graphene.String)

class TranscriptionResult(graphene.ObjectType):
    timestamps = graphene.List(Timestamp)
    # Add a field to return the full transcript text
    full_transcript_content = graphene.String()
    # Add the new analysis fields
    analysis_data = graphene.Field(AnalysisResult)
    key_insights = graphene.Field(KeyInsightsResult)
    summary_content = graphene.Field(SummaryResult) 

# Define approximate weights for each stage (add up to 100)
# This helps in calculating a smooth, continuous overall progress percentage.
STAGE_WEIGHTS = {
    'start': 0,
    'youtube_api_fetch': 10,   # Reduced
    'download': 10,            # Reduced
    'file_decode': 5,
    'document_parsing': 10,    # Reduced
    'audio_chunking': 15,      # Increased slightly
    'ai_transcription': 40,    # Starts from 40% (reduced to make space for analysis)
    'gemini_processing': 10,
    'gemini_analysis': 10,     # New stage for analysis metrics, insights, etc.
    'complete': 100
}

def calculate_overall_progress(current_stage, stage_percentage, total_sub_steps=1, completed_sub_steps=0):
    """
    Calculates the cumulative overall progress based on current stage and its internal progress.
    `stage_percentage` is the completion of the *current substep* (e.g., 50% through current chunk).
    `total_sub_steps` and `completed_sub_steps` are for stages like 'ai_transcription' that have internal iterations.
    """
    completed_weight_sum = 0

    # Sum weights of all preceding stages
    for stage, weight in STAGE_WEIGHTS.items():
        if stage == current_stage:
            break
        completed_weight_sum += weight

    current_stage_weight = STAGE_WEIGHTS.get(current_stage, 0)

    if current_stage == 'ai_transcription' and total_sub_steps > 0:
        # For AI transcription, distribute its weight across processed chunks
        progress_within_stage = (completed_sub_steps / total_sub_steps) * current_stage_weight
    else:
        # For other stages, calculate their contribution based on their percentage within their weight
        progress_within_stage = (stage_percentage / 100) * current_stage_weight

    overall_percentage = completed_weight_sum + progress_within_stage
    return min(100, max(0, int(overall_percentage))) # Clamp between 0 and 100

# Helper function to download YouTube audio (kept for potential fallback or future direct audio processing)
def download_youtube_audio(video_id: str) -> str:
    """
    Downloads the audio from a YouTube video and saves it to a temporary file.
    Returns the path to the downloaded audio file.
    """
    temp_dir = tempfile.mkdtemp()

    # Define the final desired audio file path (e.g., 'video_id.mp3')
    final_audio_file_path = os.path.join(temp_dir, f"{video_id}.mp3")

    ydl_opts = {
        'format': 'bestaudio/best',
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'mp3',
            'preferredquality': '192',
        }],
        'outtmpl': os.path.join(temp_dir, video_id),
        'quiet': True,
        'no_warnings': True,
        'noprogress': True,
        'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'referer': 'https://www.youtube.com/',
    }

    try:
        print(f"Attempting to download audio for video ID: {video_id}")
        # Emit initial download progress
        with app.app_context():
            overall_percent = calculate_overall_progress('download', 0)
            emit('progress_update', {'type': 'overall', 'status': 'Starting audio download...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            def hook(d):
                if d['status'] == 'downloading':
                    current_percent = 0
                    if 'total_bytes' in d and 'downloaded_bytes' in d and d['total_bytes'] > 0:
                        current_percent = (d['downloaded_bytes'] / d['total_bytes']) * 100
                    elif 'total_bytes_estimate' in d and 'downloaded_bytes' in d and d['total_bytes_estimate'] > 0:
                         current_percent = (d['downloaded_bytes'] / d['total_bytes_estimate']) * 100

                    overall_percent = calculate_overall_progress('download', current_percent)
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': f'Downloading audio: {current_percent:.1f}%', 'percentage': overall_percent}, namespace='/', broadcast=True)

                elif d['status'] == 'finished':
                    overall_percent = calculate_overall_progress('download', 100)
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': 'Audio download complete.', 'percentage': overall_percent}, namespace='/', broadcast=True)

            ydl.add_progress_hook(hook)

            # Use the standard YouTube URL format for yt-dlp to ensure proper handling
            standard_youtube_url = f"https://www.youtube.com/watch?v={video_id}"
            ydl.download([standard_youtube_url])

        if not os.path.exists(final_audio_file_path):
            raise RuntimeError(f"Downloaded audio file not found at expected path: {final_audio_file_path}")

        return final_audio_file_path
    except yt_dlp.utils.DownloadError as e:
        print(f"Error downloading YouTube audio: {e}")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise RuntimeError(f"Failed to download audio: {e}")
    except Exception as e:
        print(f"An unexpected error occurred during audio download: {e}")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise RuntimeError(f"An unexpected error occurred during audio download: {e}")

# Helper function to chunk audio for transcription
def chunk_audio(audio_path: str, chunk_length_ms: int = 60000) -> list:
    """
    Chunks an audio file into smaller segments.
    Returns a list of paths to the audio chunks.
    """
    try:
        audio = AudioSegment.from_file(audio_path)
        total_length_ms = len(audio)
        num_chunks = math.ceil(total_length_ms / chunk_length_ms)

        if num_chunks == 0: # Handle empty audio files gracefully
            return []

        chunk_paths = []
        base_dir = os.path.dirname(audio_path)

        for i in range(num_chunks):
            start_ms = i * chunk_length_ms
            end_ms = min((i + 1) * chunk_length_ms, total_length_ms)
            chunk = audio[start_ms:end_ms]
            chunk_file_path = os.path.join(base_dir, f"chunk_{i}.mp3")
            chunk.export(chunk_file_path, format="mp3")
            chunk_paths.append(chunk_file_path)

            overall_percent = calculate_overall_progress('audio_chunking', ((i+1)/num_chunks)*100)
            with app.app_context():
                emit('progress_update', {'type': 'overall', 'status': f'Chunking audio: {i+1}/{num_chunks} chunks processed.', 'percentage': overall_percent}, namespace='/', broadcast=True)

        return chunk_paths
    except Exception as e:
        print(f"Error chunking audio: {e}")
        raise RuntimeError(f"Error chunking audio: {e}")

# Determine the device for Faster Whisper
device = "cuda" if ctranslate2.get_cuda_device_count() > 0 else "cpu"

# Load the Faster Whisper model globally
model = None
with app.app_context():
    try:
        overall_percent = calculate_overall_progress('start', 0) # Initial 0%
        emit('progress_update', {'type': 'overall', 'status': 'Loading AI model (Faster Whisper)...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        model = WhisperModel("base", device=device, compute_type="int8")

        overall_percent = calculate_overall_progress('start', 100) # Model loaded, but overall process hasn't started yet
        emit('progress_update', {'type': 'overall', 'status': 'AI model loaded successfully.', 'percentage': overall_percent}, namespace='/', broadcast=True)
    except Exception as e:
        print(f"Error loading Faster Whisper model: {e}")
        with app.app_context():
            emit('progress_update', {'type': 'overall', 'status': f'Error loading AI model: {e}', 'percentage': 0}, namespace='/', broadcast=True)
        raise RuntimeError(f"Failed to load Faster Whisper model: {e}")

def transcribe_audio_with_faster_whisper(audio_path: str) -> (str, list):
    """
    Transcribes an audio file using the Faster Whisper model.
    Returns the full transcript and a list of timestamps with text.
    """
    if model is None:
        raise RuntimeError("Faster Whisper model not loaded.")

    try:
        segments, info = model.transcribe(audio_path, beam_size=5)
        full_transcript = ""
        timestamps = []
        for segment in segments:
            segment_text = segment.text.strip()
            if segment_text: # Ensure text is not empty
                full_transcript += segment_text + " "
                timestamps.append({"time": segment.start, "text": segment_text})
        return full_transcript.strip(), timestamps
    except Exception as e:
        print(f"Error during Faster Whisper transcription: {e}")
        raise RuntimeError(f"Transcription failed: {e}")


# Your existing API key (replace with your actual key if it's still a placeholder here)
# You should retrieve this from environment variables in a production setup
API_KEY = os.environ.get("GEMINI_API_KEY") # Prioritize environment variable
if not API_KEY:
    API_KEY = "AIzaSyCajcPzzVJCvRSBFtSwuv-Fu-Pr6lUNqBk" # Fallback

# Configurable Gemini model name (environment override for backwards/forwards compatibility)
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")

def get_gemini_response_base(chat_history_for_gemini: list, generation_config: dict = None) -> str:
    """Generic function to call Gemini API with given chat history and generation config.

    This function is defensive: it supports the candidate/content/parts format used in older Gemini
    responses, but also falls back to other possible shapes or raw text. It strips markdown fences
    and tries to return the best textual response available.
    """
    api_url = os.environ.get("GEMINI_API_URL") or f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={API_KEY}"

    payload = {
        "contents": chat_history_for_gemini,
        "generationConfig": generation_config or {
            "temperature": 0.7,
            "topK": 40,
            "topP": 0.95,
            "maxOutputTokens": 1000
        }
    }

    try:
        response = requests.post(api_url, headers={'Content-Type': 'application/json'}, data=json.dumps(payload), timeout=60)
        response.raise_for_status()
        result = None
        try:
            result = response.json()
        except Exception:
            # Fall back to raw text if response isn't JSON
            raw_text = response.text
            # Detect quota/billing style errors in raw text
            if raw_text and 'quota' in raw_text.lower():
                raise GeminiAPIError('quota_exceeded', raw_text)
            cleaned = _strip_code_fences_and_markdown(raw_text)
            return cleaned or "AI did not provide a clear response."

        # Prefer common 'candidates' -> 'content' -> 'parts' path
        try:
            if isinstance(result, dict) and result.get("candidates"):
                cand = result["candidates"][0]
                # Older shape: candidates[].content.parts[].text
                if isinstance(cand.get("content"), dict) and cand["content"].get("parts"):
                    parts = cand["content"]["parts"]
                    texts = [p.get("text", "") for p in parts if isinstance(p, dict)]
                    cleaned = "".join(texts).strip()
                    if cleaned:
                        return _strip_code_fences_and_markdown(cleaned)

                # Sometimes 'content' is just a string
                if isinstance(cand.get("content"), str) and cand.get("content"):
                    return _strip_code_fences_and_markdown(cand.get("content"))

                # Newer shapes may put text at candidates[0]['output'] or similar
                if cand.get("output"):
                    return _strip_code_fences_and_markdown(cand.get("output"))

        except Exception as e:
            print(f"Warning parsing 'candidates' shape: {e}")

        # Try some common alternative keys
        if isinstance(result, dict):
            for k in ("output", "text", "response", "content"):
                if k in result and isinstance(result[k], str) and result[k].strip():
                    return _strip_code_fences_and_markdown(result[k])

        # Fallback: inspect the raw text body
        raw_text = response.text
        cleaned = _strip_code_fences_and_markdown(raw_text)
        return cleaned or "AI did not provide a clear response."

    except requests.exceptions.HTTPError as e:
        print(f"HTTP Error during Gemini API call in get_gemini_response_base: {e}")
        try:
            print(f"Response status code: {response.status_code}")
            print(f"Response text: {response.text}")
            if response.text and 'quota' in response.text.lower():
                raise GeminiAPIError('quota_exceeded', response.text)
            error_json = response.json()
            print(f"Response JSON: {json.dumps(error_json, indent=2)}")
            if "error" in error_json and "message" in error_json["error"]:
                return f"API Error: {error_json['error']['message']}"
            else:
                return f"An HTTP error occurred: {e}"
        except GeminiAPIError:
            raise
        except Exception:
            return f"An HTTP error occurred: {e}. Non-JSON response."
    except requests.exceptions.RequestException as e:
        print(f"Network error communicating with AI: {e}")
        return f"Network error communicating with AI: {e}"
    except Exception as e:
        print(f"Unexpected error in get_gemini_response_base: {e}")
        return f"An unexpected error occurred: {e}."

def orchestrate_user_request(user_message: str, current_transcript: str = None, youtube_url: str = None) -> dict:
    """
    The main agent to understand user intent and orchestrate responses.
    This version embeds the Gemini API calls directly within the function.
    Returns a dictionary with 'action', 'content', and 'ai_chat_message'.
    """
    # Determine the context source
    if current_transcript:
        context_type = "text_content"
        context_data = current_transcript
        context_description = "the provided text content"
    elif youtube_url:
        context_type = "youtube_video_url"
        context_data = youtube_url
        context_description = f"the YouTube video at {youtube_url} (transcript not explicitly provided yet)"
    else:
        context_type = "general"
        context_data = ""
        context_description = "your request"

    intent_prompt = f"""
    You are an AI assistant for a transcription and content modification application.
    Your goal is to understand the user's request and determine if it's a general question or a command to modify content.
    The current context is related to {context_description}.

    Based on the user's message, classify the intent and provide the necessary details.

    If the user wants to:
    - Get a summary of {context_description}: Return "SUMMARY".
    - Translate {context_description}: Return "TRANSLATE: [target_language]".
    - Rephrase {context_description} (e.g., make it funny, formal, simpler): Return "REPHRASE: [rephrasing_instruction]".
    - Extract keywords from {context_description}: Return "KEYWORDS".
    - Get a story based on {context_description}: Return "STORY_CONVERSION: [story_style]".
    - Get a podcast script based on {context_description}: Return "PODCAST_SCRIPT".
    - Ask a general question about the {context_description} or the application itself, including questions about who owns or is speaking in the content: Return "GENERAL_QUESTION".
    - If the request is unclear, a greeting, or not related to the above: Return "UNSUPPORTED_COMMAND".

    Examples (assuming video content from transcript):
    User: Can you summarize this video for me?
    Output: SUMMARY

    User: Translate this video's content to Spanish.
    Output: TRANSLATE: Spanish

    User: Make the video's content funnier.
    Output: REPHRASE: make it funny and engaging

    User: Give me keywords for SEO from the video.
    Output: KEYWORDS

    User: Convert this video into a short story.
    Output: STORY_CONVERSION: short story

    User: Create a podcast script from the video's main themes.
    Output: PODCAST_SCRIPT

    User: Hi, how are you?
    Output: UNSUPPORTED_COMMAND

    User: Who is speaking in this video?
    Output: GENERAL_QUESTION

    User message: "{user_message}"
    Output:
    """

    chat_history_for_intent = [{ "role": "user", "parts": [{ "text": intent_prompt }] }]

    intent_config = {
        "temperature": 0.1,
        "topK": 1,
        "topP": 0.95,
        "maxOutputTokens": 50
    }

    intent_response = get_gemini_response_base(chat_history_for_intent, intent_config).strip()
    print(f"Gemini Intent Response: {intent_response}")

    action = "UNSUPPORTED_COMMAND"
    ai_chat_message = "I couldn't quite understand that. Can you please rephrase?"
    processed_content = ""

    content_generation_config = {
        "temperature": 0.7,
        "topK": 40,
        "topP": 0.95,
        "maxOutputTokens": 2000
    }

    # Prepare chat history for content generation, including current_transcript if available
    chat_history_for_content = []
    if current_transcript:
        chat_history_for_content.append({"role": "user", "parts": [{"text": f"Given the following content:\n\n{current_transcript}"}]})

    if intent_response == "SUMMARY":
        action = "SUMMARY"
        ai_chat_message = "Generating a summary for you. Please check the main panel."
        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 0)
            emit('progress_update', {'type': 'overall', 'status': 'Generating summary with AI...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        if current_transcript:
            prompt = "Summarize the content concisely."
            chat_history_for_content.append({"role": "user", "parts": [{"text": prompt}]})
            processed_content = get_gemini_response_base(chat_history_for_content, content_generation_config)
        else:
            processed_content = "Please provide content (transcript or YouTube video) to summarize."

        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 100) # Assuming completion for this specific gemini task
            emit('progress_update', {'type': 'overall', 'status': 'Summary generated.', 'percentage': overall_percent}, namespace='/', broadcast=True)
            emit('progress_update', {'type': 'overall', 'status': 'Processing complete!', 'percentage': 100}, namespace='/', broadcast=True)


    elif intent_response.startswith("TRANSLATE:"):
        action = "TRANSLATE"
        target_language = intent_response.split(":", 1)[1].strip()
        ai_chat_message = f"Translating your content to {target_language}. Check the main panel!"
        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 0)
            emit('progress_update', {'type': 'overall', 'status': f'Translating to {target_language} with AI...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        if current_transcript:
            prompt = f"Translate the content to {target_language}."
            chat_history_for_content.append({"role": "user", "parts": [{"text": prompt}]})
            processed_content = get_gemini_response_base(chat_history_for_content, content_generation_config)
        else:
            processed_content = "Please provide content (transcript or YouTube video) to translate."

        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 100)
            emit('progress_update', {'type': 'overall', 'status': 'Translation complete.', 'percentage': overall_percent}, namespace='/', broadcast=True)
            emit('progress_update', {'type': 'overall', 'status': 'Processing complete!', 'percentage': 100}, namespace='/', broadcast=True)


    elif intent_response.startswith("REPHRASE:"):
        action = "REPHRASE"
        rephrasing_instruction = intent_response.split(":", 1)[1].strip()
        ai_chat_message = f"Rephrasing your content based on your request. Take a look at the main panel!"
        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 0)
            emit('progress_update', {'type': 'overall', 'status': 'Rephrasing content with AI...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        if current_transcript:
            prompt = f"Rephrase the content. {rephrasing_instruction}."
            chat_history_for_content.append({"role": "user", "parts": [{"text": prompt}]})
            processed_content = get_gemini_response_base(chat_history_for_content, content_generation_config)
        else:
            processed_content = "Please provide content (transcript or YouTube video) to rephrase."

        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 100)
            emit('progress_update', {'type': 'overall', 'status': 'Rephrasing complete.', 'percentage': overall_percent}, namespace='/', broadcast=True)
            emit('progress_update', {'type': 'overall', 'status': 'Processing complete!', 'percentage': 100}, namespace='/', broadcast=True)


    elif intent_response == "KEYWORDS":
        action = "KEYWORDS"
        ai_chat_message = "Extracting keywords for you. They will appear in the main panel."
        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 0)
            emit('progress_update', {'type': 'overall', 'status': 'Extracting keywords with AI...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        if current_transcript:
            prompt = "Extract a list of relevant SEO keywords from the content. Provide them as a comma-separated list."
            chat_history_for_content.append({"role": "user", "parts": [{"text": prompt}]})
            processed_content = get_gemini_response_base(chat_history_for_content, content_generation_config)
        else:
            processed_content = "Please provide content (transcript or YouTube video) to extract keywords from."

        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 100)
            emit('progress_update', {'type': 'overall', 'status': 'Keyword extraction complete.', 'percentage': overall_percent}, namespace='/', broadcast=True)
            emit('progress_update', {'type': 'overall', 'status': 'Processing complete!', 'percentage': 100}, namespace='/', broadcast=True)


    elif intent_response.startswith("STORY_CONVERSION:"):
        action = "STORY_CONVERSION"
        story_style = intent_response.split(":", 1)[1].strip()
        ai_chat_message = f"Converting content into a {story_style} story. Check the main panel!"
        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 0)
            emit('progress_update', {'type': 'overall', 'status': f'Generating {story_style} story with AI...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        if current_transcript:
            prompt = f"Convert the content into a {story_style} story."
            chat_history_for_content.append({"role": "user", "parts": [{"text": prompt}]})
            processed_content = get_gemini_response_base(chat_history_for_content, content_generation_config)
        else:
            processed_content = "Please provide content (transcript or YouTube video) to convert into a story."

        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 100)
            emit('progress_update', {'type': 'overall', 'status': 'Story conversion complete.', 'percentage': overall_percent}, namespace='/', broadcast=True)
            emit('progress_update', {'type': 'overall', 'status': 'Processing complete!', 'percentage': 100}, namespace='/', broadcast=True)


    elif intent_response == "PODCAST_SCRIPT":
        action = "PODCAST_SCRIPT"
        ai_chat_message = "Generating a podcast script for you. Check the main panel!"
        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 0)
            emit('progress_update', {'type': 'overall', 'status': 'Generating podcast script with AI...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        if current_transcript:
            prompt = "Create a podcast script based on the main themes and content."
            chat_history_for_content.append({"role": "user", "parts": [{"text": prompt}]})
            processed_content = get_gemini_response_base(chat_history_for_content, content_generation_config)
        else:
            processed_content = "Please provide content (transcript or YouTube video) to generate a podcast script from."

        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 100)
            emit('progress_update', {'type': 'overall', 'status': 'Podcast script generated.', 'percentage': overall_percent}, namespace='/', broadcast=True)
            emit('progress_update', {'type': 'overall', 'status': 'Processing complete!', 'percentage': 100}, namespace='/', broadcast=True)


    elif intent_response == "GENERAL_QUESTION":
        action = "GENERAL_QUESTION"
        chat_history_general = []
        if current_transcript:
            chat_history_general.append({ "role": "user", "parts": [{ "text": f"Regarding the following content:\n\n{current_transcript}\n\n{user_message}" }] })
        elif youtube_url:
            chat_history_general.append({ "role": "user", "parts": [{ "text": f"Regarding the YouTube video at {youtube_url}, {user_message}" }] })
        else:
            chat_history_general.append({ "role": "user", "parts": [{ "text": user_message }] })

        ai_chat_message = get_gemini_response_base(chat_history_general) # Use the base function for general questions too
        processed_content = "" # General questions don't typically produce a 'processed_content' block
        with app.app_context():
            overall_percent = calculate_overall_progress('complete', 100) # Assuming general questions are quick and conclude overall process
            emit('progress_update', {'type': 'overall', 'status': 'General query addressed.', 'percentage': overall_percent}, namespace='/', broadcast=True)

    return {
        "action": action,
        "ai_chat_message": ai_chat_message,
        "processed_content": processed_content if action != "GENERAL_QUESTION" else ""
    }

class TranscribeVideo(graphene.Mutation):
    class Arguments:
        url = graphene.String(required=True)

    Output = TranscriptionResult

    def mutate(self, info, url):
        video_id = None
        audio_file_path = None
        temp_dir_to_clean = None

        try:
            parsed_url = urlparse(url)
            query_params = parse_qs(parsed_url.query)
            video_id = query_params.get('v', [None])[0]

            if not video_id:
                # Handle various YouTube URL formats to extract video ID
                # This part was modified to correctly extract video_id from different YouTube URLs
                if 'youtube.com' in parsed_url.hostname:
                    if parsed_url.path.startswith('/watch'):
                        video_id = query_params.get('v', [None])[0]
                    elif parsed_url.path.startswith('/embed/'):
                        video_id = parsed_url.path.split('/')[2]
                    elif parsed_url.path.startswith('/v/'):
                        video_id = parsed_url.path.split('/')[2]
                elif 'youtu.be' in parsed_url.hostname:
                    video_id = parsed_url.path[1:]

            if not video_id:
                raise ValueError("Invalid YouTube URL provided: Could not extract video ID.")

            with app.app_context():
                overall_percent = calculate_overall_progress('start', 0)
                emit('progress_update', {'type': 'overall', 'status': 'Starting transcription process...', 'percentage': overall_percent}, namespace='/', broadcast=True)

            full_transcript_text = ""
            timestamps_list = []
            transcript_found_via_api = False

            # Attempt to get transcript directly using youtube_transcript_api first
            try:
                overall_percent = calculate_overall_progress('youtube_api_fetch', 10)
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Attempting to fetch YouTube transcript via API...', 'percentage': overall_percent}, namespace='/', broadcast=True)

                fetched_transcript_list = None
                try:
                    fetched_transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'], preserve_formatting=True)
                    print("English transcript fetched using youtube_transcript_api.")
                    overall_percent = calculate_overall_progress('youtube_api_fetch', 50)
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': 'English transcript fetched successfully.', 'percentage': overall_percent}, namespace='/', broadcast=True)
                except NoTranscriptFound:
                    print("English transcript not found. Attempting to fetch any available transcript.")
                    overall_percent = calculate_overall_progress('youtube_api_fetch', 20)
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': 'English transcript not found. Searching for any available language.', 'percentage': overall_percent}, namespace='/', broadcast=True)

                    available_transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
                    found_any_transcript_obj = False
                    for t_obj in available_transcripts:
                        try:
                            fetched_transcript_list = t_obj.fetch()
                            print(f"Found and fetched transcript in language: {t_obj.language_code}")
                            overall_percent = calculate_overall_progress('youtube_api_fetch', 80)
                            with app.app_context():
                                emit('progress_update', {'type': 'overall', 'status': f'Transcript fetched successfully in {t_obj.language_code}.', 'percentage': overall_percent}, namespace='/', broadcast=True)
                            found_any_transcript_obj = True
                            break
                        except Exception as fetch_err:
                            print(f"Could not fetch transcript for {t_obj.language_code}: {fetch_err}")
                            continue

                    if not found_any_transcript_obj:
                        raise NoTranscriptFound("No transcripts found in any language for this video.")

                if fetched_transcript_list:
                    processed_timestamps = []
                    full_transcript_parts = []

                    for t in fetched_transcript_list:
                        if isinstance(t, dict) and 'start' in t and 'text' in t:
                            segment_start = t['start']
                            segment_text = t['text']
                        elif hasattr(t, 'start') and hasattr(t, 'text'):
                            segment_start = t.start
                            segment_text = t.text
                        else:
                            print(f"Warning: Unexpected transcript segment format: {t}")
                            continue
                        segment_text_stripped = segment_text.strip()
                        if segment_text_stripped:
                            full_transcript_parts.append(segment_text_stripped)
                            processed_timestamps.append({"time": segment_start, "text": segment_text_stripped})

                    if processed_timestamps:
                        full_transcript_text = " ".join(full_transcript_parts)
                        timestamps_list = processed_timestamps
                        transcript_found_via_api = True
                    else:
                        print("Warning: Processed transcript content is empty despite successful fetch.")
                        raise RuntimeError("API returned empty transcript after processing.")
                else:
                    print(f"Warning: Fetched transcript content is empty or has an unhandled structure after fetching: {type(fetched_transcript_list)}")
                    raise RuntimeError("API returned empty or malformed transcript.")

            except (NoTranscriptFound, VideoUnavailable, TranscriptsDisabled, RuntimeError) as e:
                print(f"Failed to fetch YouTube transcript directly ({e}). Falling back to AI transcription...")
                overall_percent = calculate_overall_progress('youtube_api_fetch', 100) # Mark API fetch as 'complete' (failed)
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': f'Direct transcript unavailable ({e}). Falling back to AI transcription.', 'percentage': overall_percent}, namespace='/', broadcast=True)
                transcript_found_via_api = False


            if not transcript_found_via_api:
                if model is None:
                    raise RuntimeError("Faster Whisper model not loaded. Cannot perform AI transcription.")

                # Download audio progress handled within download_youtube_audio
                audio_file_path = download_youtube_audio(video_id)
                print(f"DEBUG: Downloaded audio to: {audio_file_path}")
                temp_dir_to_clean = os.path.dirname(audio_file_path)

                # Chunk audio progress handled within chunk_audio
                chunk_paths = chunk_audio(audio_file_path)
                print(f"DEBUG: Chunked audio into {len(chunk_paths)} chunks.")
                total_chunks = len(chunk_paths)

                full_transcript_parts = []
                aggregated_timestamps = []

                for i, chunk_path in enumerate(chunk_paths):
                    overall_percent = calculate_overall_progress('ai_transcription', 0, total_chunks, i) # 0% for current chunk
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': f'Transcribing chunk {i+1}/{total_chunks} with AI model...', 'percentage': overall_percent}, namespace='/', broadcast=True)

                    chunk_transcript, chunk_timestamps = transcribe_audio_with_faster_whisper(chunk_path)
                    full_transcript_parts.append(chunk_transcript)

                    chunk_start_time_ms = i * 60000
                    for ts in chunk_timestamps:
                        adjusted_time = ts['time'] + (chunk_start_time_ms / 1000)
                        aggregated_timestamps.append({"time": adjusted_time, "text": ts['text']})

                    if os.path.exists(chunk_path):
                        os.remove(chunk_path)

                full_transcript_text = " ".join(full_transcript_parts)
                timestamps_list = aggregated_timestamps

                overall_percent = calculate_overall_progress('ai_transcription', 100, total_chunks, total_chunks) # Ensure transcription stage is 100%
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'AI transcription complete.', 'percentage': overall_percent}, namespace='/', broadcast=True)
            
            analysis_data = None
            key_insights_data = None
            summary_data = None

            if full_transcript_text:
                # We no longer call external Gemini for analysis to avoid quota-related failures.
                # Instead, deterministically produce top terms from the transcript for Quick Insights.
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Generating quick insights...', 'percentage': calculate_overall_progress('gemini_analysis', 50)}, namespace='/', broadcast=True)

                top_terms = get_top_terms_from_transcript(full_transcript_text, top_n=10)
                key_insights_data = {'title': 'Top Terms', 'points': top_terms}

                # Structured summary is still useful and produced via Gemini (kept as-is)
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Generating structured summary...', 'percentage': calculate_overall_progress('gemini_analysis', 95)}, namespace='/', broadcast=True)
                summary_data = get_gemini_summary(full_transcript_text)

                # analysis_data intentionally omitted (removed analysis tab/server-side analysis)
                analysis_data = None

                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Quick insights ready.', 'percentage': calculate_overall_progress('gemini_analysis', 100)}, namespace='/', broadcast=True)

            # Final overall complete signal
            overall_percent = calculate_overall_progress('complete', 100)
            with app.app_context():
                emit('progress_update', {'type': 'overall', 'status': 'Processing complete!', 'percentage': overall_percent}, namespace='/', broadcast=True)

            return TranscriptionResult(
                timestamps=timestamps_list,
                full_transcript_content=full_transcript_text,
                analysis_data=analysis_data,
                key_insights=KeyInsightsResult(**key_insights_data) if key_insights_data else None,
                # MODIFIED: Ensure SummaryResult is created from the generated data
                summary_content=SummaryResult(**summary_data) if summary_data and summary_data.get('title') else None
            )

        except Exception as e:
            print(f"Error in transcribeVideo mutation: {e}")
            with app.app_context():
                emit('progress_update', {'type': 'overall', 'status': f'Error: {e}', 'percentage': 0}, namespace='/', broadcast=True) # Reset progress on error
            raise e
        finally:
            if temp_dir_to_clean and os.path.exists(temp_dir_to_clean):
                print(f"Cleaning up temporary directory: {temp_dir_to_clean}")
                shutil.rmtree(temp_dir_to_clean)

class TranscribeFile(graphene.Mutation):
    class Arguments:
        file_content = graphene.String(required=True)
        file_name = graphene.String(required=True)
        file_mime_type = graphene.String(required=True)

    Output = TranscriptionResult

    def mutate(self, info, file_content, file_name, file_mime_type):
        temp_dir_to_clean = None
        temp_file_path = None
        try:
            overall_percent = calculate_overall_progress('start', 0)
            with app.app_context():
                emit('progress_update', {'type': 'overall', 'status': 'Starting file processing...', 'percentage': overall_percent}, namespace='/', broadcast=True)
                overall_percent = calculate_overall_progress('file_decode', 0)
                emit('progress_update', {'type': 'overall', 'status': 'Decoding file content...', 'percentage': overall_percent}, namespace='/', broadcast=True)

            temp_dir_to_clean = tempfile.mkdtemp()
            if ',' in file_content:
                header, base64_data = file_content.split(',', 1)
            else:
                base64_data = file_content

            decoded_file_content = base64.b64decode(base64_data)

            temp_file_path = os.path.join(temp_dir_to_clean, file_name)

            with open(temp_file_path, 'wb') as f:
                f.write(decoded_file_content)

            overall_percent = calculate_overall_progress('file_decode', 100)
            with app.app_context():
                emit('progress_update', {'type': 'overall', 'status': 'File decoded and saved temporarily.', 'percentage': overall_percent}, namespace='/', broadcast=True)

            full_transcript_text = ""
            timestamps_list = []

            if file_mime_type == 'application/pdf':
                overall_percent = calculate_overall_progress('document_parsing', 0)
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Extracting text from PDF...', 'percentage': overall_percent}, namespace='/', broadcast=True)
                try:
                    with open(temp_file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        total_pages = len(reader.pages)
                        for page_num in range(total_pages):
                            full_transcript_text += reader.pages[page_num].extract_text() or ""
                            current_page_percent = ((page_num + 1) / total_pages) * 100
                            overall_percent = calculate_overall_progress('document_parsing', current_page_percent)
                            with app.app_context():
                                emit('progress_update', {'type': 'overall', 'status': f'Extracting text from PDF: Page {page_num+1}/{total_pages}', 'percentage': overall_percent}, namespace='/', broadcast=True)

                    overall_percent = calculate_overall_progress('document_parsing', 100)
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': 'Text extracted from PDF.', 'percentage': overall_percent}, namespace='/', broadcast=True)
                except Exception as e:
                    raise RuntimeError(f"Error extracting text from PDF: {e}")
            elif file_mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or \
                 file_mime_type == 'application/msword': # .docx and .doc
                overall_percent = calculate_overall_progress('document_parsing', 0)
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Extracting text from DOCX...', 'percentage': overall_percent}, namespace='/', broadcast=True)
                try:
                    doc = Document(temp_file_path)
                    total_paragraphs = len(doc.paragraphs)
                    for i, para in enumerate(doc.paragraphs):
                        full_transcript_text += para.text + "\n"
                        current_para_percent = ((i + 1) / total_paragraphs) * 100
                        overall_percent = calculate_overall_progress('document_parsing', current_para_percent)
                        with app.app_context():
                            emit('progress_update', {'type': 'overall', 'status': f'Extracting text from DOCX: Paragraph {i+1}/{total_paragraphs}', 'percentage': overall_percent}, namespace='/', broadcast=True)

                    overall_percent = calculate_overall_progress('document_parsing', 100)
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': 'Text extracted from DOCX.', 'percentage': overall_percent}, namespace='/', broadcast=True)
                except Exception as e:
                    raise RuntimeError(f"Error extracting text from DOCX: {e}")
            elif file_mime_type == 'text/plain':
                overall_percent = calculate_overall_progress('document_parsing', 0)
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Reading text from TXT...', 'percentage': overall_percent}, namespace='/', broadcast=True)
                try:
                    with open(temp_file_path, 'r', encoding='utf-8') as file:
                        full_transcript_text = file.read()
                    overall_percent = calculate_overall_progress('document_parsing', 100)
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': 'Text read from TXT.', 'percentage': overall_percent}, namespace='/', broadcast=True)
                except Exception as e:
                    raise RuntimeError(f"Error reading text from TXT: {e}")
            elif file_mime_type.startswith('audio/') or file_mime_type.startswith('video/'):
                overall_percent = calculate_overall_progress('audio_chunking', 0)
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Chunking audio from file...', 'percentage': overall_percent}, namespace='/', broadcast=True)

                chunk_paths = chunk_audio(temp_file_path)
                total_chunks = len(chunk_paths)

                full_transcript_parts = []
                aggregated_timestamps = []

                for i, chunk_path in enumerate(chunk_paths):
                    overall_percent = calculate_overall_progress('ai_transcription', 0, total_chunks, i) # 0% for current chunk, based on overall AI transcr.
                    with app.app_context():
                        emit('progress_update', {'type': 'overall', 'status': f'Transcribing chunk {i+1}/{total_chunks} with AI model...', 'percentage': overall_percent}, namespace='/', broadcast=True)

                    chunk_transcript, chunk_timestamps = transcribe_audio_with_faster_whisper(chunk_path)
                    full_transcript_parts.append(chunk_transcript)

                    chunk_start_time_seconds = i * 60
                    for ts in chunk_timestamps:
                        adjusted_time = ts['time'] + chunk_start_time_seconds
                        aggregated_timestamps.append({"time": adjusted_time, "text": ts['text']})

                    if os.path.exists(chunk_path):
                        os.remove(chunk_path)

                full_transcript_text = " ".join(full_transcript_parts)
                timestamps_list = aggregated_timestamps

                overall_percent = calculate_overall_progress('ai_transcription', 100, total_chunks, total_chunks) # Ensure transcription stage is 100%
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'AI transcription complete.', 'percentage': overall_percent}, namespace='/', broadcast=True)

            else:
                raise ValueError(f"Unsupported file MIME type: {file_mime_type}")

            analysis_data = None
            key_insights_data = None
            summary_data = None

            if full_transcript_text:
                with app.app_context():
                    # Update progress for starting analysis
                    emit('progress_update', {'type': 'overall', 'status': 'Starting AI-powered analysis...', 'percentage': calculate_overall_progress('gemini_analysis', 0)}, namespace='/', broadcast=True)
                    
                    # Generate metrics
                    emit('progress_update', {'type': 'overall', 'status': 'Generating analysis metrics...', 'percentage': calculate_overall_progress('gemini_analysis', 20)}, namespace='/', broadcast=True)
                metrics = get_gemini_analysis_metrics(full_transcript_text)

                with app.app_context():
                    # Generate rephrase suggestions
                    emit('progress_update', {'type': 'overall', 'status': 'Generating rephrasing suggestions...', 'percentage': calculate_overall_progress('gemini_analysis', 50)}, namespace='/', broadcast=True)
                rephrase_suggestions = get_gemini_rephrase_suggestions(full_transcript_text)

                with app.app_context():
                    # Generate engagement data
                    emit('progress_update', {'type': 'overall', 'status': 'Simulating engagement data...', 'percentage': calculate_overall_progress('gemini_analysis', 80)}, namespace='/', broadcast=True)
                engagement_chart_data = get_gemini_engagement_chart_data(full_transcript_text)

                with app.app_context():
                    # Generate key insights
                    emit('progress_update', {'type': 'overall', 'status': 'Generating key insights...', 'percentage': calculate_overall_progress('gemini_analysis', 90)}, namespace='/', broadcast=True)
                key_insights_data = get_gemini_key_insights(full_transcript_text)
                
                # MODIFIED: Call get_gemini_summary to get structured summary data
                with app.app_context():
                    emit('progress_update', {'type': 'overall', 'status': 'Generating structured summary...', 'percentage': calculate_overall_progress('gemini_analysis', 95)}, namespace='/', broadcast=True)
                summary_data = get_gemini_summary(full_transcript_text)
                
                # For documents, use Gemini for summarization as an initial action
                if file_mime_type.startswith('application/') or file_mime_type.startswith('text/'):
                    summary = get_gemini_response_base([{"role": "user", "parts": [{"text": f"Summarize the following text concisely:\n\n{full_transcript_text}"}]}], {
                        "temperature": 0.7,
                        "topK": 40,
                        "topP": 0.95,
                        "maxOutputTokens": 500
                    })
                    timestamps_list = [{"time": 0.0, "text": summary}]
                
                # Combine into the AnalysisResult object
                analysis_data = AnalysisResult(
                    metrics=metrics,
                    rephraseSuggestions=rephrase_suggestions,
                    engagementChartData=engagement_chart_data
                )
                
                with app.app_context():
                    # Mark analysis stage as complete
                    emit('progress_update', {'type': 'overall', 'status': 'AI analysis complete!', 'percentage': calculate_overall_progress('gemini_analysis', 100)}, namespace='/', broadcast=True)

            # Final overall complete signal
            overall_percent = calculate_overall_progress('complete', 100)
            with app.app_context():
                emit('progress_update', {'type': 'overall', 'status': 'Processing complete!', 'percentage': overall_percent}, namespace='/', broadcast=True)

            return TranscriptionResult(
                timestamps=timestamps_list,
                full_transcript_content=full_transcript_text,
                analysis_data=analysis_data,
                key_insights=KeyInsightsResult(**key_insights_data) if key_insights_data else None,
                summary_content=SummaryResult(**summary_data) if summary_data and summary_data.get('title') else None
            )

        except Exception as e:
            print(f"Error in transcribeFile mutation: {e}")
            with app.app_context():
                emit('progress_update', {'type': 'overall', 'status': f'Error: {e}', 'percentage': 0}, namespace='/', broadcast=True)
            raise e
        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                print(f"Cleaned up original temporary file: {temp_file_path}")
            if temp_dir_to_clean and os.path.exists(temp_dir_to_clean):
                print(f"Cleaning up temporary directory: {temp_dir_to_clean}")
                shutil.rmtree(temp_dir_to_clean)


class ChatResponse(graphene.ObjectType):
    ai_chat_message = graphene.String()
    processed_content = graphene.String()

class ProcessChatCommand(graphene.Mutation):
    class Arguments:
        user_message = graphene.String(required=True)
        current_transcript = graphene.String()
        youtube_url = graphene.String()

    Output = ChatResponse

    def mutate(self, info, user_message, current_transcript=None, youtube_url=None):
        print(f"Received chat command: '{user_message}'")
        if current_transcript:
            print(f"With transcript length: {len(current_transcript)}")
        if youtube_url:
            print(f"With YouTube URL: {youtube_url}")

        with app.app_context():
            overall_percent = calculate_overall_progress('gemini_processing', 0)
            emit('progress_update', {'type': 'overall', 'status': 'Analyzing your request with AI...', 'percentage': overall_percent}, namespace='/', broadcast=True)

        try:
            orchestration_result = orchestrate_user_request(user_message, current_transcript, youtube_url)
            ai_chat_message = orchestration_result["ai_chat_message"]
            processed_content = orchestration_result["processed_content"]
            action_taken = orchestration_result["action"]

            if action_taken == "GENERAL_QUESTION":
                processed_content = ""
                with app.app_context():
                    overall_percent = calculate_overall_progress('complete', 100) # Assuming general questions are quick and conclude overall process
                    emit('progress_update', {'type': 'overall', 'status': 'General query addressed.', 'percentage': overall_percent}, namespace='/', broadcast=True)
            else:
                 with app.app_context():
                    overall_percent = calculate_overall_progress('complete', 100) # Assuming the specific content action is complete
                    emit('progress_update', {'type': 'overall', 'status': f'Content {action_taken.lower()} completed!', 'percentage': overall_percent}, namespace='/', broadcast=True)

        except Exception as e:
            print(f"Error orchestrating chat command: {e}")
            ai_chat_message = f"I apologize, but I encountered an error while processing your request: {e}. Please try again or rephrase."
            processed_content = current_transcript if current_transcript else ""
            with app.app_context():
                emit('progress_update', {'type': 'overall', 'status': f'Error processing request: {e}', 'percentage': 0}, namespace='/', broadcast=True)

        return ChatResponse(ai_chat_message=ai_chat_message, processed_content=processed_content)


def analyze_text_helper(transcript_text: str) -> dict:
    """Analyze a piece of transcript text and return normalized analysis, insights, and summary.

    Returns a dict with keys: analysis_data (dict), key_insights (dict), summary_data (dict)
    """
    # Use Graphene types for analysis_data for consistent behavior
    analysis_data = AnalysisResult(metrics=[], rephraseSuggestions=[], engagementChartData=[])
    key_insights_data = {'title': 'No Insights Available', 'points': []}
    summary_data = {'title': 'No Summary Available', 'subheader': '', 'points': []}

    if not transcript_text or not transcript_text.strip():
        return {'analysis_data': analysis_data, 'key_insights': key_insights_data, 'summary_data': summary_data}

    try:
        with app.app_context():
            emit('progress_update', {'type': 'overall', 'status': 'Starting AI-powered analysis...', 'percentage': calculate_overall_progress('gemini_analysis', 0)}, namespace='/', broadcast=True)
        # Produce deterministic quick insights (top terms) locally to avoid external API issues
        top_terms = get_top_terms_from_transcript(transcript_text, top_n=10)
        key_insights_data = {'title': 'Top Terms', 'points': top_terms}

        # Keep summary generation if it exists, but it's optional
        try:
            summary_data = get_gemini_summary(transcript_text)
        except Exception:
            summary_data = {'title': 'No Summary Available', 'subheader': '', 'points': []}

        # analysis_data intentionally omitted
        analysis_data = None

        with app.app_context():
            emit('progress_update', {'type': 'overall', 'status': 'Quick insights ready.', 'percentage': calculate_overall_progress('gemini_analysis', 100)}, namespace='/', broadcast=True)

        return {
            'analysisData': analysis_data,
            'keyInsights': key_insights_data,
            'summaryContent': summary_data
        }

    except Exception as e:
        print(f"Error during analyze_text_helper: {e}")
        _save_raw_gemini_response('analyze_text_error', str(e))
        return {'analysis_data': analysis_data, 'key_insights': key_insights_data, 'summary_data': summary_data}


class AnalyzeText(graphene.Mutation):
    class Arguments:
        text = graphene.String(required=True)

    analysisData = graphene.Field(AnalysisResult)
    keyInsights = graphene.Field(KeyInsightsResult)
    summaryContent = graphene.Field(SummaryResult)

    def mutate(self, info, text):
        try:
            result = analyze_text_helper(text)
            return AnalyzeText(
                analysisData=result['analysisData'],
                keyInsights=KeyInsightsResult(**result['keyInsights']) if result.get('keyInsights') else None,
                summaryContent=SummaryResult(**result['summaryContent']) if result.get('summaryContent') else None
            )
        except Exception as e:
            print(f"Error in AnalyzeText mutation: {e}")
            return AnalyzeText(analysis_data=None, key_insights=None, summary_data=None)


# New Gemini functions for analysis
def get_gemini_analysis_metrics(transcript_text: str) -> list:
    """Sends transcript to Gemini to get analysis metrics."""
    if not transcript_text.strip():
        return []

    prompt = """Return strictly valid JSON only, without markdown or extra explanation.
    Analyze the following transcript content for performance metrics.
    Provide the output as a JSON array of objects for four key areas, where each object has 'title', 'value', 'change', and 'changeType'.
    'changeType' should be 'positive', 'negative', or 'neutral'.
    For 'Filler Word Count', 'positive' means lower count.
    Example:
    [
        {"title": "Trend Score", "value": "88", "change": "+5.2%", "changeType": "positive"},
        {"title": "Audience Engagement", "value": "92%", "change": "+8%", "changeType": "positive"},
        {"title": "Clarity Score", "value": "95%", "change": "+1.5%", "changeType": "positive"},
        {"title": "Filler Word Count", "value": "12", "change": "-3", "changeType": "positive"}
    ]
    Transcript:
    """ + transcript_text

    chat_history = [{"role": "user", "parts": [{"text": prompt}]}]
    try:
        response_text = get_gemini_response_base(chat_history, generation_config={
            "temperature": 0.5,
            "maxOutputTokens": 1200
        })

        parsed = _extract_first_json(response_text)
        if parsed is None:
            print("Error: Could not parse JSON from Gemini analysis metrics response.")
            print(f"Raw response: {response_text}")
            _save_raw_gemini_response('analysis_metrics', response_text)
            return []

        if isinstance(parsed, list):
            return parsed
        else:
            print(f"Warning: Expected list from analysis metrics, got {type(parsed)}. Returning empty list.")
            return []
    except GeminiAPIError as ge:
        print(f"Gemini API error in analysis metrics: {ge}")
        _save_raw_gemini_response('analysis_metrics_error', str(ge))
        return []

    # Try to extract JSON payload safely from the response
    parsed = _extract_first_json(response_text)
    if parsed is None:
        print("Error: Could not parse JSON from Gemini analysis metrics response.")
        print(f"Raw response: {response_text}")
        _save_raw_gemini_response('analysis_metrics', response_text)
        return []

    normalized = []
    if isinstance(parsed, list):
        for item in parsed:
            try:
                title = str(item.get('title', 'Unknown')) if isinstance(item, dict) else str(item)
                value = str(item.get('value', '')) if isinstance(item, dict) else ''
                change = str(item.get('change', '')) if isinstance(item, dict) else ''
                changeType = str(item.get('changeType', 'neutral')) if isinstance(item, dict) else 'neutral'
                # sanitize changeType
                if changeType not in ('positive', 'negative', 'neutral'):
                    changeType = 'neutral'
                normalized.append({
                    'title': title,
                    'value': value,
                    'change': change,
                    'changeType': changeType
                })
            except Exception as e:
                print(f"Skipping malformed metric item: {e}")
        return normalized
    else:
        print(f"Warning: Expected list from analysis metrics, got {type(parsed)}. Returning empty list.")
        return []

def get_gemini_rephrase_suggestions(transcript_text: str) -> list:
    """Sends transcript to Gemini to get rephrasing suggestions."""
    if not transcript_text.strip():
        return []

    prompt = """Return strictly valid JSON only, without markdown or extra explanation.
    Given the following transcript, identify 3-5 sentences that could be rephrased for better clarity, conciseness, or engagement.
    Provide the output as a JSON array of objects, where each object has 'original' and 'suggested'.
    Example:
    [
        {"original": "The integration of smart grids is essential for managing diverse energy sources.", "suggested": "Smart grids are key to handling different types of energy."}
    ]
    Transcript:
    """ + transcript_text

    chat_history = [{"role": "user", "parts": [{"text": prompt}]}]
    try:
        response_text = get_gemini_response_base(chat_history, generation_config={
            "temperature": 0.7,
            "maxOutputTokens": 1600
        })

        parsed = _extract_first_json(response_text)
        if parsed is None:
            print("Error: Could not parse JSON from Gemini rephrase suggestions response.")
            print(f"Raw response: {response_text}")
            _save_raw_gemini_response('rephrase_suggestions', response_text)
            return []
    except GeminiAPIError as ge:
        print(f"Gemini API error in rephrase suggestions: {ge}")
        _save_raw_gemini_response('rephrase_suggestions_error', str(ge))
        return []

    normalized = []
    if isinstance(parsed, list):
        for item in parsed:
            if isinstance(item, dict):
                original = str(item.get('original', '')).strip()
                suggested = str(item.get('suggested', '')).strip()
                if original and suggested:
                    normalized.append({'original': original, 'suggested': suggested})
            else:
                # If item is string, skip or convert
                continue
        return normalized
    else:
        print(f"Warning: Expected list from rephrase suggestions, got {type(parsed)}. Returning empty list.")
        return []


import json
import re

def get_gemini_engagement_chart_data(transcript_text: str) -> list:
    """Sends transcript to Gemini to get simulated engagement chart data."""
    if not transcript_text.strip():
        return []

    prompt = """Return strictly valid JSON only, without markdown or extra explanation.
    Simulate audience engagement percentages for a presentation based on its transcript content, broken down into 10% segments.
    Consider factors like complexity, rhetorical questions, calls to action, narrative flow, and potential for audience interest.
    Provide the output as a JSON array of objects, where each object has 'segment' (e.g., '0%', '10%', ...) and 'engagement' (an integer from 0-100).
    Ensure there are exactly 11 segments (0% to 100%).
    Example:
    [
        {"segment": "0%", "engagement": 65},
        {"segment": "10%", "engagement": 72}
    ]
    Transcript:
    """ + transcript_text

    chat_history = [{"role": "user", "parts": [{"text": prompt}]}]
    try:
        response_text = get_gemini_response_base(chat_history, generation_config={
            "temperature": 0.8, # Higher temperature for more varied simulation
            "maxOutputTokens": 800
        })

        if not response_text or not response_text.strip():
            print("Warning: Gemini returned an empty response for engagement data. Returning default data.")
            return [{"segment": f"{i*10}%", "engagement": 50} for i in range(11)]

        parsed = _extract_first_json(response_text)
        if parsed is None or not isinstance(parsed, list):
            print("Warning: Could not parse valid engagement JSON; returning default segments. Raw response below:")
            print(response_text)
            _save_raw_gemini_response('engagement_data', response_text)
            return [{"segment": f"{i*10}%", "engagement": 50} for i in range(11)]
    except GeminiAPIError as ge:
        print(f"Gemini API error in engagement data: {ge}")
        _save_raw_gemini_response('engagement_data_error', str(ge))
        return [{"segment": f"{i*10}%", "engagement": 50} for i in range(11)]

    # Ensure exactly 11 segments, fill missing or truncate extras
    if len(parsed) != 11:
        print("Warning: Gemini did not return exactly 11 engagement segments. Attempting to adjust.")
        segments_map = {item['segment']: item['engagement'] for item in parsed if 'segment' in item and 'engagement' in item}
        corrected_data = []
        for i in range(11):
            segment_label = f"{i*10}%"
            engagement_value = segments_map.get(segment_label, 50) # Default to 50 if missing
            corrected_data.append({"segment": segment_label, "engagement": engagement_value})
        return corrected_data

    return parsed
        
def get_gemini_key_insights(transcript_text: str) -> dict:
    """Sends transcript to Gemini to get key insights, formatted as a title and a list of bullet points."""
    if not transcript_text.strip():
        return {"title": "No Insights Available", "points": []}

    prompt = """Return strictly valid JSON only, without markdown or extra explanation.
    Analyze the following transcript and identify the 3-7 most important and representative keywords or key phrases.
    These should be the core terms or concepts that best describe the main topics or critical elements discussed in the transcript. Do not add any extra explanations, headlines, or descriptive sentences beyond the keywords/phrases themselves.

    Return the output as a JSON object with 'title' (string, e.g., "Key Transcript Keywords") and 'points' (array of strings, where each string is a keyword or key phrase).

    Example:
    {
        "title": "Key Transcript Keywords",
        "points": [
            "TCS",
            "Codevita",
            "Placements",
            "Gemma",
            "MCP",
            "War",
            "Vivobook S15"
        ]
    }
    Transcript:
    """ + transcript_text

    chat_history = [{"role": "user", "parts": [{"text": prompt}]}]
    try:
        chat_history = [{"role": "user", "parts": [{"text": prompt}]}]
        response_text = get_gemini_response_base(chat_history, generation_config={
            "temperature": 0.2, # Lower temperature for more factual, keyword extraction
            "maxOutputTokens": 300 # Increased to reduce truncation
        })

        parsed = _extract_first_json(response_text)
        if parsed is None or not isinstance(parsed, dict) or "title" not in parsed or "points" not in parsed:
            print("Warning: Gemini insights response malformed. Expected 'title' and 'points'.")
            print(f"Raw response: {response_text}")
            _save_raw_gemini_response('key_insights', response_text)
            return {"title": "Insights Generation Failed", "points": ["Could not parse insights from AI."]}

        # Normalize to ensure title (string) and points (list of strings)
        normalized = {
            'title': str(parsed.get('title', 'Insights')) if isinstance(parsed, dict) else 'Insights',
            'points': [str(p) for p in parsed.get('points', [])] if isinstance(parsed, dict) and isinstance(parsed.get('points', []), list) else []
        }
        return normalized
    except GeminiAPIError as ge:
        print(f"Gemini API error in key insights: {ge}")
        _save_raw_gemini_response('key_insights_error', str(ge))
        return {"title": "AI Quota Exceeded", "points": [str(ge)]}

def get_gemini_summary(transcript_text: str) -> dict:
    """Sends transcript to Gemini to get a structured summary."""
    if not transcript_text.strip():
        return {"title": "No Summary Available", "subheader": "", "points": []}

    prompt = """Return strictly valid JSON only, without markdown or extra explanation.
    Create a detailed summary of the following transcript.
    The summary should be structured as a JSON object with a main "title", a concise "subheader", and a "points" array containing 3-5 key bullet points.
    Focus on the main themes, key arguments, and significant conclusions.
    The output should be clean JSON, without any markdown formatting like ```json.
    Example:
    {
        "title": "Key Takeaways from the Presentation",
        "subheader": "A high-level overview of the future of renewable energy.",
        "points": [
            "Solar panel efficiency is improving due to new materials.",
            "Offshore wind farms present significant potential.",
            "AI-powered smart grids are crucial for energy distribution."
        ]
    }
    Transcript:
    """ + transcript_text

    try:
        chat_history = [{"role": "user", "parts": [{"text": prompt}]}]
        response_text = get_gemini_response_base(chat_history, generation_config={
            "temperature": 0.5,
            "maxOutputTokens": 1200
        })

        print(f"Gemini summary response: {response_text}")

        parsed = _extract_first_json(response_text)
        if parsed is None or not isinstance(parsed, dict):
            print("Error decoding Gemini summary JSON: could not extract JSON object.")
            print(f"Raw response: {response_text}")
            _save_raw_gemini_response('summary', response_text)
            return {"title": "Summary Generation Failed", "subheader": "Could not parse summary from AI.", "points": []}

        # Normalize summary shape
        normalized = {
            'title': str(parsed.get('title', 'Summary Generation Failed')) if isinstance(parsed, dict) else 'Summary Generation Failed',
            'subheader': str(parsed.get('subheader', '')) if isinstance(parsed, dict) else '',
            'points': [str(p) for p in parsed.get('points', [])] if isinstance(parsed, dict) and isinstance(parsed.get('points', []), list) else []
        }
        return normalized
    except GeminiAPIError as ge:
        print(f"Gemini API error in summary: {ge}")
        _save_raw_gemini_response('summary_error', str(ge))
        return {"title": "AI Quota Exceeded", "subheader": str(ge), "points": []}


def get_top_terms_from_transcript(transcript_text: str, top_n: int = 10) -> list:
    """A simple, deterministic extractor that returns the top N meaningful terms from the transcript.

    This avoids calling external APIs and is resilient to malformed AI responses.
    It lowercases the text, strips punctuation, removes common stopwords and short tokens, and returns
    the most frequent terms (ties broken by lexical order).
    """
    if not transcript_text or not transcript_text.strip():
        return []

    # Minimal stopword set — keep it small to avoid pulling extra deps
    stopwords = {
        'the','and','is','in','to','a','of','it','that','this','for','on','with','as','are','was','but','be','by','an','or','from','at','we','you','i','they','he','she','them','his','her'
    }

    # Normalize and tokenize
    text = transcript_text.lower()
    # Replace non-alphanumeric characters with spaces
    text = re.sub(r"[^a-z0-9\s]", ' ', text)
    tokens = [t for t in text.split() if len(t) >= 3 and not t.isdigit() and t not in stopwords]
    if not tokens:
        return []

    # Count frequencies
    freqs = {}
    for t in tokens:
        freqs[t] = freqs.get(t, 0) + 1

    # Sort by frequency desc, then alphabetically
    sorted_terms = sorted(freqs.items(), key=lambda kv: (-kv[1], kv[0]))
    top_terms = [term for term, _ in sorted_terms[:top_n]]
    return top_terms

class Query(graphene.ObjectType):
    hello = graphene.String(name=graphene.String(default_value="World"))

    def resolve_hello(self, info, name):
        return f"Hello {name}"

class Mutation(graphene.ObjectType):
    transcribe_video = TranscribeVideo.Field()
    transcribe_file = TranscribeFile.Field()
    process_chat_command = ProcessChatCommand.Field()
    analyze_text = AnalyzeText.Field()

schema = graphene.Schema(query=Query, mutation=Mutation)

app.add_url_rule(
    '/graphql',
    view_func=GraphQLView.as_view(
        'graphql',
        schema=schema,
        graphiql=True
    )
)


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    print(f"Running Flask app with SocketIO on port {port}")
    # Set allow_unsafe_werkzeug=True only for development, disable in production
    socketio.run(app, debug=True, host='0.0.0.0', port=port, use_reloader=False, allow_unsafe_werkzeug=True)